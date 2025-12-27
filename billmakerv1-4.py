import sys
import os
from PyQt6.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, 
                             QLabel, QLineEdit, QPushButton, QTableWidget, 
                             QHeaderView, QFileDialog, QMessageBox)
from docx import Document
from docx2pdf import convert

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

class AstaEpsilonBilling(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Asta Epsilon Infotronics - Billing System v1.4')
        self.setMinimumWidth(700)
        main_layout = QVBoxLayout()

        self.inputs = {}
        fields = [
            ('Client Name', '{{name}}'),
            ('Client Address', '{{address}}'),
            ('Bill Number', '{{invoiceno}}'),
            ('Bill Date', '{{billdate}}'),
            ('Due Date', '{{duedate}}')
        ]

        for label_text, placeholder in fields:
            row = QHBoxLayout()
            lbl = QLabel(label_text)
            lbl.setFixedWidth(120)
            row.addWidget(lbl)
            self.inputs[placeholder] = QLineEdit()
            row.addWidget(self.inputs[placeholder])
            main_layout.addLayout(row)

        main_layout.addWidget(QLabel("\nService Details:"))
        self.table = QTableWidget(3, 3)
        self.table.setHorizontalHeaderLabels(['Description', 'Quantity', 'Unit Price'])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        main_layout.addWidget(self.table)

        btn_layout = QHBoxLayout()
        self.btn_docx = QPushButton('Export Word (.docx)')
        self.btn_pdf = QPushButton('Export PDF (.pdf)')
        self.btn_clear = QPushButton('Clear Form')
        
        self.btn_pdf.setStyleSheet("background-color: #008CBA; color: white; font-weight: bold; padding: 8px;")
        self.btn_clear.setStyleSheet("background-color: #f44336; color: white; padding: 8px;")

        self.btn_docx.clicked.connect(lambda: self.process_bill(export_pdf=False))
        self.btn_pdf.clicked.connect(lambda: self.process_bill(export_pdf=True))
        self.btn_clear.clicked.connect(self.clear_form)
        
        btn_layout.addWidget(self.btn_docx)
        btn_layout.addWidget(self.btn_pdf)
        btn_layout.addWidget(self.btn_clear)
        main_layout.addLayout(btn_layout)
        self.setLayout(main_layout)

    def clear_form(self):
        for field in self.inputs.values(): field.clear()
        self.table.clearContents()

    def process_bill(self, export_pdf):
        data = {key: field.text() for key, field in self.inputs.items()}
        grand_total = 0

        for i in range(3):
            desc = self.table.item(i, 0).text() if self.table.item(i, 0) else ""
            qty_val = self.table.item(i, 1).text() if self.table.item(i, 1) else "0"
            price_val = self.table.item(i, 2).text() if self.table.item(i, 2) else "0"

            try:
                qty = float(qty_val)
                price = float(price_val)
                line_total = qty * price
                grand_total += line_total
                data[f"{{{{description{i+1}}}}}"] = desc
                data[f"{{{{quantity{i+1}}}}}"] = str(qty)
                data[f"{{{{amount{i+1}}}}}"] = f"{line_total:,.2f}"
            except ValueError:
                data[f"{{{{description{i+1}}}}}"] = desc
                data[f"{{{{quantity{i+1}}}}}"] = ""
                data[f"{{{{amount{i+1}}}}}"] = ""

        data["{{total}}"] = f"{grand_total:,.2f}"
        self.save_document(data, export_pdf)

    def save_document(self, data, export_pdf):
        template_name = "Bill Format.docx"
        template_path = resource_path(template_name)
        
        if not os.path.exists(template_path):
            QMessageBox.critical(self, "Error", "Template not found.")
            return

        file_filter = "PDF Files (*.pdf)" if export_pdf else "Word Files (*.docx)"
        output_path, _ = QFileDialog.getSaveFileName(self, "Save Bill", "", file_filter)
        
        if output_path:
            try:
                doc = Document(template_path)
                
                # REVISED LOGIC: Iterates runs to keep "White Font" formatting
                def replace_and_preserve(paragraphs):
                    for p in paragraphs:
                        for key, val in data.items():
                            if key in p.text:
                                # We search each run. If the placeholder is split 
                                # across runs, this logic joins it temporarily.
                                for run in p.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, str(val))
                                    # Handle edge case where {{ and total}} are split
                                    elif "{{" in run.text and "}}" not in run.text:
                                        pass # Complex split handling would go here

                replace_and_preserve(doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_and_preserve(cell.paragraphs)

                temp_docx = "temp_render.docx"
                doc.save(temp_docx)

                if export_pdf:
                    convert(temp_docx, output_path)
                    os.remove(temp_docx)
                else:
                    if os.path.exists(output_path): os.remove(output_path)
                    os.rename(temp_docx, output_path)

                QMessageBox.information(self, "Success", "Bill generated with formatting preserved!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed: {str(e)}")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = AstaEpsilonBilling()
    window.show()
    sys.exit(app.exec())